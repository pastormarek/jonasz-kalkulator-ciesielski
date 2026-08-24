"""
Generator formularza .docx do konsultacji ciesielskiej — pierwsza tura.

Dokument zawiera te same 50 założeń, co strona `docs/zalozenia.body.html`,
ale w formie do wypełnienia: pod każdym punktem jest pole, w które można wpisać
odpowiedź w Wordzie, Word Online albo Dokumentach Google.

Formatowanie siedzi w `docx_common.py`, wspólnym dla kolejnych tur pytań.

Uruchomienie:
    python tools/formularz-docx.py

Wynik trafia do docs/Jonasz-formularz-zalozenia.docx
"""

from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt  # noqa: E402

from docx_common import (  # noqa: E402
    AKCENT,
    OBRAMOWANIE,
    TLO_POLA,
    TUSZ,
    TUSZ_SLABY,
    akapit,
    cieniuj,
    naglowek,
    nowy_dokument,
    obramuj,
    pole_odpowiedzi,
    wypisz_punkty,
)

# --- treść merytoryczna ------------------------------------------------------
# Krotka: (numer, założenie, pytanie albo None)

SEKCJE = [
    (
        "A. Wymiary i geometria połaci",
        "Podstawa wszystkiego. Jeśli tu jest błąd, reszta też jest.",
        [
            (
                1,
                "Rozpiętość mierzymy w poprzek budynku, po zewnętrznych krawędziach murłat. "
                "Od tej krawędzi liczymy okap i tam wypada zacios.",
                "Czy tak się to podaje, czy raczej w osiach ścian albo po licu muru?",
            ),
            (
                2,
                "Okap podaje się jako wysunięcie w poziomie poza krawędź murłaty. Wzdłuż krokwi "
                "wychodzi z tego więcej — przy 35° z 60 cm robi się 73 cm.",
                None,
            ),
            (
                3,
                "Długość krokwi = (połowa rozpiętości + okap) ÷ cos(kąt). Zakładamy, że krokiew "
                "jest cięta pionowo na obu końcach — w kalenicy i przy okapie.",
                None,
            ),
            (
                4,
                "Do każdej sztuki doliczamy naddatek na docięcie — domyślnie 5 cm.",
                "Ile realnie zostawiasz zapasu na krokwi?",
            ),
            (
                5,
                "Wysokość kalenicy = połowa rozpiętości × tg(kąt), liczona od poziomu górnej "
                "krawędzi murłaty — nie od stropu ani od podłogi poddasza.",
                None,
            ),
            (
                6,
                "W więźbie krokwiowej krokwie stykają się w kalenicy czołowo, dokładnie w osi. "
                "Nie odejmujemy nic na grubość drugiej krokwi.",
                "Czy robisz styk czołowy, na zakładkę, czy na wzajemne wcięcie?",
            ),
            (
                7,
                "Powierzchnia połaci = powierzchnia rzutu ÷ cos(kąt). Rzut liczymy razem "
                "z okapami i wysunięciem szczytowym.",
                None,
            ),
            (
                8,
                "Kąt można podać w stopniach albo jako spadek w procentach — jedno przelicza się "
                "na drugie. 100% to 45°.",
                None,
            ),
        ],
    ),
    (
        "B. Rozstaw krokwi",
        None,
        [
            (
                9,
                "Rozstaw liczymy w osiach krokwi. Podajesz największy dopuszczalny, a program "
                "dzieli długość budynku na tyle równych pól, żeby go nie przekroczyć. Przy "
                "budynku 12 m i maksimum 90 cm wychodzi 14 pól po 85,1 cm.",
                None,
            ),
            (
                10,
                "Skrajne krokwie licują z krawędziami ścian szczytowych, więc ich osie są "
                "cofnięte o pół grubości krokwi do środka.",
                "Czy skrajna krokiew idzie przy samej ścianie, czy odsuwa się ją od szczytu?",
            ),
            (
                11,
                "Obok rozstawu w osiach pokazujemy prześwit — ile jest między bokami sąsiednich "
                "krokwi. To ta liczba jest istotna przy wełnie.",
                None,
            ),
        ],
    ),
    (
        "C. Zacios na murłacie",
        "Tu najbardziej zależy nam na sprawdzeniu, bo to liczby odmierzane wprost na drewnie.",
        [
            (
                12,
                "Liczymy zacios siodłowy: płaszczyzna pozioma leży na murłacie, pionowa opiera "
                "się o jej bok.",
                "Czy taki wykonujesz najczęściej, czy jednak inny rodzaj wcięcia?",
            ),
            (
                13,
                "Głębokość zaciosu mierzymy prostopadle do krokwi, nie w pionie. Domyślnie 3 cm.",
                None,
            ),
            (
                14,
                "Siodło = głębokość ÷ sin(kąt), pięta = głębokość ÷ cos(kąt). Przy dachu 35° "
                "i zaciosie 3 cm daje to siodło 5,2 cm i piętę 3,7 cm.",
                "Czy te wymiary zgadzają się z tym, co odmierzasz na krokwi?",
            ),
            (
                15,
                "Program ostrzega, gdy zacios przekracza 1/3 wysokości krokwi. Przy krokwi 18 cm "
                "granica wypada na 6 cm.",
                None,
            ),
            (
                16,
                "Ostrzegamy też, gdy siodło nie mieści się na szerokości murłaty — przy małych "
                "kątach robi się bardzo długie.",
                None,
            ),
        ],
    ),
    (
        "D. Jętki",
        None,
        [
            (
                17,
                "Rozpiętość jętki = rozpiętość − 2 × wysokość ÷ tg(kąt).",
                None,
            ),
            (
                18,
                "Do zamówienia doliczamy 2 × grubość krokwi, zakładając, że jętka jest przybijana "
                "z boku krokwi na zakładkę.",
                "Robisz to na zakładkę z boku, czy wcinasz jętkę w krokiew? Jeśli wcinasz, "
                "długość powinna być inna.",
            ),
            (
                19,
                "Liczymy jedną jętkę na każdą parę krokwi.",
                "Czy zdarza się dawać co drugą parę?",
            ),
            (
                20,
                "Wysokość jętki mierzymy od poziomu murłaty do jej dolnej krawędzi.",
                None,
            ),
        ],
    ),
    (
        "E. Więźba płatwiowo-kleszczowa",
        "Najsłabiej opracowana część. Tu spodziewamy się najwięcej poprawek.",
        [
            (
                21,
                "Rozstaw słupów działa tak samo jak rozstaw krokwi: dzielimy długość budynku na "
                "równe pola, nie przekraczając zadanego maksimum.",
                None,
            ),
            (
                22,
                "Wysokość słupa jest ZGADYWANA jako połowa wysokości kalenicy. Program nie wie, "
                "na jakim poziomie jest strop ani jak wysoka jest ścianka kolankowa, więc nie ma "
                "z czego jej policzyć.",
                "O co program powinien zapytać, żeby policzyć to porządnie: o wysokość ścianki "
                "kolankowej, o poziom stropu, czy po prostu wprost o długość słupa?",
            ),
            (
                23,
                "Kleszcze liczymy jako parę desek o długości równej rozpiętości plus 2 × grubość "
                "krokwi.",
                "Czy kleszcze wychodzą poza krokwie, a jeśli tak, to o ile?",
            ),
            (
                24,
                "Miecz traktujemy jako przekątną trójkąta o równych ramionach, cięty pod 45°. "
                "Przy ramieniu 80 cm wychodzi 113 cm.",
                None,
            ),
        ],
    ),
    (
        "F. Dach kopertowy",
        "Kąty sprawdziliśmy z tablicami — dla dachu 45° wychodzi znane 35°16′ na kulawce "
        "i 30° na sfazowaniu krożyny.",
        [
            (
                25,
                "Zakładamy równe spadki wszystkich czterech połaci, więc naroże w rzucie biegnie "
                "dokładnie pod 45°.",
                "Czy zdarzają się koperty o różnych spadkach na szczytach?",
            ),
            (
                26,
                "Krożyna jest zawsze łagodniejsza od połaci: tg(krożyna) = tg(połaci) ÷ √2. "
                "Dla dachu 35° daje to 26,3°.",
                None,
            ),
            (
                27,
                "Kulawki skracają się równomiernie, o rozstaw podzielony przez cos(kąt).",
                None,
            ),
            (
                28,
                "Ukos kulawki podajemy jako arcus tangens cosinusa kąta połaci. Dla 35° to 39,3°, "
                "dla 45° — 35,3°.",
                "Czy to jest kąt, który realnie nastawiasz na pile, czy podajesz go inaczej?",
            ),
            (
                29,
                "Sfazowanie grzbietu krożyny = arcus tangens sinusa kąta krożyny. Dla dachu 35° "
                "wychodzi 23,9°.",
                None,
            ),
            (
                30,
                "Kulawek każdego rozmiaru liczymy osiem sztuk — cztery naroża, przy każdym "
                "kulawki z dwóch stron.",
                None,
            ),
            (
                31,
                "Kalenica jest krótsza od budynku o rozpiętość. Gdy budynek jest kwadratowy, "
                "kalenica schodzi do punktu i wychodzi dach namiotowy.",
                None,
            ),
        ],
    ),
    (
        "G. Drewno, długości i łączenie",
        None,
        [
            (
                32,
                "Rozróżniamy drewno z półki (3, 4, 5 i 6 m) oraz cięte na wymiar w tartaku, do 12 m.",
                "Czy te długości się zgadzają? Do ilu metrów realnie da się zamówić belkę i czy "
                "jest granica, powyżej której nikt tego nie dowiezie?",
            ),
            (
                33,
                "Murłaty i płatwie dzielimy automatycznie na kawałki mieszczące się w belce "
                "handlowej — leżą na podporze całą długością. Murłata 12 m to dwie sztuki po 6 m.",
                None,
            ),
            (
                34,
                "Krokwi i jętek nigdy nie dzielimy same z siebie. Jeśli nie mieszczą się "
                "w dostępnej belce, program mówi o tym wprost i podpowiada: zamówić na wymiar "
                "albo świadomie włączyć łączenie.",
                None,
            ),
            (
                35,
                "Krokiew wolno złożyć z dwóch kawałków, ale styk musi wypaść nad podporą — ścianą "
                "kolankową, wieńcem albo płatwią. Podaje się odległość podpory od murłaty, "
                "a program liczy oba odcinki.",
                "Czy tak się to robi? Czy są sytuacje, w których i tak nie wolno łączyć?",
            ),
            (
                36,
                "Przyjęliśmy nakładkę 60 cm na styku.",
                "Ile robi się realnie i czym się to spina — śrubami, gwoździami, płytką "
                "kolczastą? Czy „nakładka” to w ogóle właściwe słowo?",
            ),
            (
                37,
                "Plan cięcia dobiera belki tak, żeby zostało jak najmniej odpadu. Jętka 1,93 m "
                "nie idzie z trzymetrówki, tylko po dwie z czterometrówki — odpad spada z 36% "
                "na 10%.",
                None,
            ),
            (
                38,
                "Rzaz piły liczymy jako 4 mm na każde cięcie.",
                "Czy to sensowna wartość, czy przy tej skali w ogóle nie ma znaczenia?",
            ),
        ],
    ),
    (
        "H. Łacenie i warstwy",
        None,
        [
            (
                39,
                "Łaty liczymy jako długość połaci podzieloną przez rozstaw, plus jeden rząd "
                "zapasu, razy szerokość połaci.",
                "Czy łata okapowa i kalenicowa wymagają osobnego potraktowania? Czy pierwszy "
                "rozstaw przy okapie jest inny niż reszta?",
            ),
            (
                40,
                "Kontrłaty — jedna na każdej krokwi, na całej jej długości.",
                None,
            ),
            (
                41,
                "Membrana: powierzchnia połaci plus 15% na zakłady i docinki.",
                "Czy 15% wystarcza?",
            ),
            (
                42,
                "Ocieplenie liczymy tylko na to, co między krokwiami — odejmujemy udział drewna "
                "w połaci.",
                None,
            ),
            (
                43,
                "Deskowanie: powierzchnia połaci plus 10% na docinki.",
                None,
            ),
        ],
    ),
    (
        "I. Łączniki i impregnat",
        "Cała ta sekcja to nasze przypuszczenia. Każda liczba jest do zakwestionowania.",
        [
            (
                44,
                "Kotwy mocujące murłatę do wieńca — co 1,5 m.",
                "Jaki rozstaw stosujesz i czym kotwisz?",
            ),
            (
                45,
                "Kątowniki ciesielskie — dwa na każde oparcie krokwi, po dziesięć wkrętów "
                "na kątownik.",
                "Czy dajesz kątowniki na każdą krokiew, czy co którąś?",
            ),
            (
                46,
                "Połączenie w kalenicy — cztery gwoździe albo wkręty na parę krokwi.",
                "Ile realnie i czym?",
            ),
            (
                47,
                "Jętka — dwie śruby M12 z podkładką na każdy koniec.",
                "Śruby czy gwoździe? Ile sztuk?",
            ),
            (
                48,
                "Impregnat — 0,2 litra na metr kwadratowy powierzchni drewna, w dwóch warstwach. "
                "Dla przykładowego dachu wychodzi 95 litrów, co wydaje nam się dużo.",
                "Ile impregnatu schodzi realnie na taki dach?",
            ),
        ],
    ),
    (
        "J. Komin i okno dachowe",
        None,
        [
            (
                49,
                "Otwór przerywa tyle krokwi, ile ich osi wpada w jego szerokość. Przerwana "
                "krokiew zamienia się w dwa krótsze odcinki — nad i pod otworem.",
                "Czy komin da się zwykle „ominąć” rozstawem, zamiast przecinać krokiew?",
            ),
            (
                50,
                "Na każdy otwór liczymy dwa wymiany — nad i pod nim — o długości szerokości "
                "otworu plus 2 × grubość krokwi.",
                "Czy wymian opiera się na sąsiednich krokwiach, czy wchodzi w nie wcięciem?",
            ),
        ],
    ),
]

PRZYKLAD = [
    ("Długość krokwi z okapem", "5,62 m"),
    ("Sama połać, bez okapu", "4,88 m"),
    ("Wysokość kalenicy nad murłatą", "2,80 m"),
    ("Rozstaw krokwi w osiach", "851 mm"),
    ("Prześwit między krokwiami", "771 mm"),
    ("Krokwi razem", "30 szt."),
    ("Powierzchnia połaci", "143,8 m²"),
    ("Zacios — siodło", "52 mm"),
    ("Zacios — pięta", "37 mm"),
    ("Jętka, długość do zamówienia", "1,93 m"),
    ("Krokwie 80 × 180 mm", "30 × 6 m"),
    ("Murłaty 140 × 140 mm", "4 × 6 m"),
    ("Jętki 80 × 160 mm", "8 × 4 m"),
    ("Drewno do kupienia", "3,47 m³"),
    ("Łaty", "511 mb"),
    ("Membrana z zakładami", "165 m²"),
    ("Kątowniki ciesielskie", "60 szt."),
    ("Impregnat", "95 l"),
]

NIE_ROBI = [
    "Nie sprawdza nośności i nie dobiera przekrojów — to musi wyjść z projektu konstrukcyjnego.",
    "Nie liczy obciążenia śniegiem ani wiatrem.",
    "Nie zna lukarn, naczółków ani wolego oka.",
    "Nie liczy obróbek blacharskich, rynien i wyłazów.",
    "Nie zna cen — pokazuje ilości, nie koszt.",
]


def zbuduj():
    dok = nowy_dokument()

    # ---------- strona tytułowa ----------
    akapit(dok, "KONSULTACJA CIESIELSKA", 9, bold=True, kolor=AKCENT, po=4)
    akapit(dok, "Kalkulator więźby — założenia do sprawdzenia", 22, bold=True, po=10)

    akapit(
        dok,
        "Marek buduje aplikację, która liczy więźbę dachową: długość i rozstaw krokwi, "
        "zaciosy, naroża dachu kopertowego oraz zestawienie drewna do zakupu. Liczby zgadzają "
        "się matematycznie — ale matematyka nie wie, jak się to robi na budowie.",
        po=8,
    )
    akapit(
        dok,
        "Poniżej jest wszystko, co program przyjmuje za pewnik. Pod każdym punktem jest pole "
        "na odpowiedź — wystarczy kliknąć i pisać. Nie trzeba wypełniać wszystkiego.",
        po=8,
    )
    akapit(
        dok,
        "Punkty z pytaniem wprost są wyróżnione kolorem. To miejsca, w których zgadywaliśmy "
        "i gdzie pomoc jest najbardziej potrzebna.",
        po=14,
    )

    # metryczka
    meta = dok.add_table(rows=1, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, etykieta in enumerate(("Wypełnia", "Data")):
        kom = meta.cell(0, i)
        cieniuj(kom, TLO_POLA)
        obramuj(kom, OBRAMOWANIE)
        p = kom.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(etykieta)
        r.font.size = Pt(8)
        r.font.bold = True
        r.font.color.rgb = TUSZ_SLABY
        pusty = kom.add_paragraph()
        pusty.paragraph_format.space_after = Pt(0)
        pusty.add_run("").font.size = Pt(11)

    akapit(dok, po=10)

    # ---------- skrót dla zabieganych ----------
    naglowek(dok, "Jeśli masz mało czasu", 13, przed=10, po=6)
    akapit(
        dok,
        "Najbardziej zależy nam na tych punktach. Odpowiedź na same te miejsca i tak będzie "
        "ogromną pomocą:",
        po=6,
    )
    for tekst in (
        "1 — od czego mierzy się rozpiętość",
        "12–16 — zacios: jak głęboko i jak wymierzasz",
        "22 — wysokość słupów, gdy jest ścianka kolankowa",
        "35–36 — łączenie krokwi i długość nakładki",
        "39 — łaty przy okapie i przy kalenicy",
        "44–47 — łączniki: kotwy, kątowniki, gwoździe, śruby",
    ):
        akapit(dok, f"•  {tekst}", 11, po=2, wciecie=0.4)

    akapit(dok, po=8)
    akapit(
        dok,
        "Jeśli coś jest dobrze — nie trzeba nic pisać. Puste pole czytamy jako „zgadza się”.",
        10,
        kolor=TUSZ_SLABY,
        kursywa=True,
        po=6,
    )

    dok.add_page_break()

    # ---------- właściwe punkty ----------
    wypisz_punkty(dok, SEKCJE)

    # ---------- przykładowy dach ----------
    dok.add_page_break()
    naglowek(dok, "Przykładowy dach — do sprawdzenia liczb")
    akapit(
        dok,
        "Dwuspadowy, więźba krokwiowo-jętkowa. Budynek 8,00 × 12,00 m, nachylenie 35°, "
        "okap 60 cm. Krokwie 8 × 18 cm, murłata 14 × 14 cm, jętki 8 × 16 cm na wysokości "
        "2,20 m. Zacios 3 cm. Dachówka ceramiczna, łaty co 32 cm. Drewno z półki.",
        10,
        kolor=TUSZ_SLABY,
        po=10,
    )

    tab = dok.add_table(rows=len(PRZYKLAD) + 1, cols=2)
    tab.alignment = WD_TABLE_ALIGNMENT.LEFT

    naglowki = ("Pozycja", "Wynik")
    for i, tekst in enumerate(naglowki):
        kom = tab.cell(0, i)
        cieniuj(kom, TLO_POLA)
        obramuj(kom, OBRAMOWANIE)
        p = kom.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(tekst)
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = TUSZ_SLABY
        if i == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for w, (nazwa, wartosc) in enumerate(PRZYKLAD, start=1):
        for i, tekst in enumerate((nazwa, wartosc)):
            kom = tab.cell(w, i)
            obramuj(kom, OBRAMOWANIE, grubosc=4)
            p = kom.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(tekst)
            r.font.size = Pt(10)
            r.font.name = "Calibri"
            r.font.color.rgb = TUSZ
            if i == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r.font.bold = True

    akapit(dok, po=8)
    akapit(
        dok,
        "Gdyby ten sam budynek przykryć dachem kopertowym: krożyna 26,3° i 7,26 m długości, "
        "ukos kulawki 39,3°, sfazowanie grzbietu 23,9°, kalenica 4,00 m.",
        10,
        kolor=TUSZ_SLABY,
        po=10,
    )
    pole_odpowiedzi(dok, "CZY TE LICZBY WYGLĄDAJĄ SENSOWNIE?", 3, pytanie=True)

    # ---------- czego nie robi ----------
    naglowek(dok, "Czego kalkulator świadomie nie robi")
    for tekst in NIE_ROBI:
        akapit(dok, f"×  {tekst}", 11, kolor=TUSZ_SLABY, po=3, wciecie=0.2)
    akapit(dok, po=6)
    pole_odpowiedzi(
        dok,
        "CZY BEZ KTÓREJŚ Z TYCH RZECZY NARZĘDZIE JEST BEZUŻYTECZNE?",
        3,
        pytanie=True,
    )

    # ---------- pole otwarte ----------
    naglowek(dok, "Czego nie zapytaliśmy, a powinniśmy")
    akapit(
        dok,
        "Najbardziej przydatne jest to, czego nie ma w żadnej tablicy: jak się to robi "
        "naprawdę, co się pomija, a co potem wychodzi na budowie.",
        po=8,
    )
    pole_odpowiedzi(dok, "CO NAM UMKNĘŁO", 8, pytanie=True)

    return dok


if __name__ == "__main__":
    wyjscie = Path(__file__).resolve().parent.parent / "docs" / "Jonasz-formularz-zalozenia.docx"
    wyjscie.parent.mkdir(parents=True, exist_ok=True)
    zbuduj().save(wyjscie)
    print(f"Zapisano: {wyjscie}")
