"""
Czwarta tura pytań do konsultacji ciesielskiej.

Numeracja jest ciągła z poprzednimi turami i zaczyna się od 87 (tura 1: 1–50,
tura 2: 51–73, tura 3: 74–86). Trzecia tura nie miała własnego formularza —
Jonasz przysłał wtedy listę wytycznych i rysunki odniesienia, stąd luka
w numeracji plików w `tools/`.

Pytania wzięły się z dwóch miejsc: z rzeczy, które trzeba było zgadnąć
przy wdrażaniu punktów 54–56 i 67, oraz z tych punktów tury 2 i 3, których
nie da się zaprogramować bez konkretnej liczby.

Uruchomienie:
    python tools/formularz4-docx.py
"""

from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx.shared import Pt  # noqa: E402

from docx_common import (  # noqa: E402
    SYGNAL,
    TUSZ_SLABY,
    akapit,
    naglowek,
    nowy_dokument,
    pole_odpowiedzi,
    wypisz_punkty,
)

# --- co zmieniły odpowiedzi z drugiej i trzeciej tury ----------------------

ZMIENIONE = [
    ("Kalenica", "tylko cięcie czołowe", "do wyboru zakładka ciesielska na pół grubości"),
    ("Długość krokwi", "do osi kalenicy", "przy zakładce dłuższa o h ÷ sin 2α"),
    ("Deska podrynnowa", "nieliczona", "osobne pole wysokości, w zestawieniu z 15% naddatku"),
    ("Cięcie przy okapie", "brak", "pionowe, 2 cm niżej niż wysokość deski"),
    ("Model 3D", "krokwie kończyły się w osi", "przy zakładce mijają się w kalenicy"),
    ("Meble", "nie było", "katalog 33 projektów z instrukcją montażu"),
]

# --- pytania ---------------------------------------------------------------

SEKCJE = [
    (
        "U. Zamek w kalenicy — domyka punkty 54, 55 i 78",
        "Zakładka jest już policzona i działa. Z Twojego opisu — „dolna krawędź krokwi "
        "dochodzi aż do górnej krawędzi kolejnej krokwi w szczycie” — wychodzi, że przy "
        "krokwi 18 cm i kącie 42° każda krokiew jest dłuższa o 18,1 cm. Zostały cztery "
        "rzeczy, których z rysunku nie odczytaliśmy na pewno.",
        [
            (
                87,
                "Wymiary zamka z Twojego schematu.",
                "Podałeś 6,0 cm, 5,4 cm i 8,0 cm przy 42°. Od czego są mierzone? Czy "
                "któraś z nich to głębokość wybrania, a któraś długość nakładki? Liczymy "
                "dziś głębokość jako połowę grubości krokwi — czy to się zgadza?",
            ),
            (
                88,
                "Czy zakładka to domyślny sposób.",
                "Zostawiliśmy wybór: cięcie czołowe albo zakładka, a domyślnie czołowe. "
                "Czy u Ciebie zakładka jest standardem i powinna być domyślna, czy "
                "faktycznie zależy to od projektu?",
            ),
            (
                89,
                "Jak ścięty jest sam koniec krokwi przy zakładce.",
                "Pionowo, czy równolegle do boku krokwi przeciwnej? Pytamy, bo to widać "
                "na rysunku 3D i chcemy narysować to poprawnie.",
            ),
            (
                90,
                "Czym spina się krokwie w zakładce.",
                "Wkręty, śruba przelotowa, czy gwoździe? Ile sztuk na jedno połączenie "
                "i jakiej długości? Dziś nie liczymy tego wcale.",
            ),
            (
                91,
                "Czy przy zakładce potrzebna jest płatew kalenicowa.",
                "Czy krokwie spięte zakładką trzymają się same, czy i tak kładzie się pod "
                "nie płatew? Od tego zależy, czy przy tym wyborze dokładać ją do "
                "zestawienia.",
            ),
        ],
    ),
    (
        "W. Okap i deska podrynnowa — domyka punkty 56 i 67",
        "Deska podrynnowa jest już osobnym polem, a krokiew tnie się pionowo 2 cm niżej "
        "niż jej wysokość — dokładnie jak w Twoim przykładzie z deską 20 cm.",
        [
            (
                92,
                "Czy te 2 cm są zawsze takie same.",
                "Czy odsadzka zależy od rodzaju rynny albo pokrycia, czy zawsze wynosi "
                "2 cm? Wpisaliśmy ją na sztywno.",
            ),
            (
                93,
                "Grubość deski podrynnowej.",
                "Przyjęliśmy deskę calową, 25 mm. Czy to właściwa grubość, czy bierzesz "
                "grubszą?",
            ),
            (
                94,
                "Gdzie ta deska biegnie.",
                "Liczymy ją tylko wzdłuż okapów. Czy idzie też po szczytach, przy "
                "wiatrownicy?",
            ),
            (
                95,
                "Cięcie poziome przy okapie.",
                "W pierwszej turze wspominałeś o cięciu pionowym ORAZ poziomym pod 90° do "
                "niego. Robimy dziś tylko pionowe. Czy poziome jest potrzebne, a jeśli "
                "tak — jak głębokie?",
            ),
        ],
    ),
    (
        "Y. Rzeczy z drugiej tury, które nadal czekają",
        "To punkty, na które już odpowiedziałeś, ale żeby je zaprogramować, brakuje nam "
        "jednej konkretnej liczby albo rozstrzygnięcia. Wypisujemy tylko te, przy których "
        "zgadywanie byłoby ryzykowne.",
        [
            (
                96,
                "Od czego liczyć długość słupa (punkty 51–53).",
                "Napisałeś, że przy dwuspadowych i kopertowych liczy się od podłogi "
                "poddasza, jeśli nie ma ściany nośnej. Czy program ma pytać osobno "
                "o poziom, na którym słup stoi, czy zawsze przyjmować podłogę poddasza?",
            ),
            (
                97,
                "Prześwit pod okno dachowe (punkt 57).",
                "Podałeś 1–2 cm z każdej strony. Czy ten sam luz zostawia się nad oknem "
                "i pod nim, czy tylko po bokach, między krokwiami?",
            ),
            (
                98,
                "Druga droga podawania rozpiętości (punkt 60).",
                "Napisałeś, że pytać należy o rozstaw zewnętrzny między murłatami, "
                "a przy obrysie budynku — dodatkowo o grubość ścianki. Czy warto dokładać "
                "tę drugą drogę, czy w praktyce wszyscy i tak podają rozstaw murłat?",
            ),
            (
                99,
                "Wiatrownice (punkt 66).",
                "Napisałeś, że to zazwyczaj łaty, po obu stronach. Jakiego przekroju "
                "i gdzie dokładnie się je przybija — na krokwiach od spodu, czy na "
                "łatach od góry?",
            ),
            (
                100,
                "Pas okapowy z podwójnej łaty (punkt 65).",
                "Druga łata odsunięta o 1–2 cm — czy jest tej samej grubości co reszta, "
                "i czy idzie na całej długości okapu?",
            ),
        ],
    ),
    (
        "Z. Rysunki i wizualizacje — z Twoich wytycznych",
        "Bierzemy się teraz za to, o co prosiłeś w ostatniej turze: strukturę drewna, "
        "podpisy elementów, pokrycie w kolorach i rysunki 2D do druku. Żeby nie zgadywać, "
        "pytamy o konkrety.",
        [
            (
                101,
                "Które elementy podpisać na modelu.",
                "Na Twoim rysunku podpisane są: płatew kalenicowa, kleszcze, krokiew "
                "zwykła, słupek i murłata. Czy to jest właściwy zestaw, czy dołożyć "
                "jeszcze coś — jętkę, miecz, kulawkę, krożynę?",
            ),
            (
                102,
                "Pokrycia na wizualizacji.",
                "Które materiały są realnie potrzebne do podglądu: dachówka ceramiczna, "
                "betonowa, blachodachówka, blacha trapezowa, blacha na rąbek, gont? "
                "I jakie kolory najczęściej wybierają klienci?",
            ),
            (
                103,
                "Łaty na modelu.",
                "Napisałeś, żeby ich nie rysować. Usunąć je z rysunku zupełnie, czy "
                "zostawić możliwość włączenia — na wypadek, gdyby ktoś chciał zobaczyć "
                "rozstaw?",
            ),
            (
                104,
                "Co musi być na rzucie 2D do druku.",
                "Prosiłeś o projekt 2D z wymiarowaniem, do zabrania na budowę. Co ma być "
                "na takiej kartce, żeby była użyteczna: rzut z góry z rozstawem krokwi, "
                "przekrój poprzeczny, detal zaciosu, tabela długości? Co jest "
                "najważniejsze?",
            ),
            (
                105,
                "Format wydruku.",
                "A4 czy A3? I czy rysunek ma być w skali, czy wystarczy, że wymiary są "
                "opisane liczbami?",
            ),
        ],
    ),
    (
        "Ź. Meble — do punktu 79",
        "Napisałeś, że meble są niedoprecyzowane i że detale są ważne. Katalog ma dziś "
        "33 projekty: ławki, stoły, donice, grządki, kompostownik, drewutnia, budki, buda "
        "dla psa, łóżko piętrowe. Każdy ma listę części, instrukcję montażu krok po kroku "
        "i rozpiskę drewna. Zanim zaczniemy dopracowywać, wolimy wiedzieć, co poprawić "
        "najpierw.",
        [
            (
                106,
                "Od czego zacząć.",
                "Które meble z tej listy są najczęściej robione i od których zacząć "
                "dopracowywanie detali?",
            ),
            (
                107,
                "Jakie złącza w meblach ogrodowych.",
                "Liczymy dziś wszystko na wkręty. Czy w meblach, które robisz, idą kołki, "
                "czopy albo śruby — a jeśli tak, to gdzie?",
            ),
            (
                108,
                "Czego brakuje w opisie mebla.",
                "Co powinno się znaleźć obok listy części i instrukcji, żeby majsterkowicz "
                "faktycznie to zrobił: rysunek złącza, kolejność cięcia, szablon "
                "w skali 1:1?",
            ),
        ],
    ),
]


def zbuduj():
    dok = nowy_dokument()

    naglowek(dok, "Kalkulator ciesielski — czwarta tura pytań", 20, przed=0, po=6)
    akapit(
        dok,
        "Dzięki za wytyczne i rysunki. Zakładka w kalenicy i deska podrynnowa są już "
        "policzone i działają — poniżej najpierw krótko, co się przez to zmieniło, "
        "a potem pytania.",
        11,
        po=14,
    )

    naglowek(dok, "Co się zmieniło po Twoich odpowiedziach")
    for nazwa, bylo, jest in ZMIENIONE:
        p = dok.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)

        r = p.add_run(f"{nazwa}: ")
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.name = "Calibri"

        r = p.add_run(f"{bylo} → ")
        r.font.size = Pt(10)
        r.font.color.rgb = TUSZ_SLABY
        r.font.name = "Calibri"

        r = p.add_run(jest)
        r.font.size = Pt(10)
        r.font.color.rgb = SYGNAL
        r.font.name = "Calibri"

    akapit(dok, po=10)
    akapit(
        dok,
        "Wydłużenie krokwi przy zakładce liczymy ze wzoru wysokość ÷ sin(2 × kąt). "
        "Wyszedł wprost z Twojego zdania o dolnej i górnej krawędzi — dla krokwi 18 cm "
        "przy 42° daje 18,1 cm.",
        10,
        kolor=TUSZ_SLABY,
        kursywa=True,
        po=6,
    )

    dok.add_page_break()

    wypisz_punkty(dok, SEKCJE)

    # --- pole otwarte ---
    naglowek(dok, "Cokolwiek jeszcze")
    akapit(
        dok,
        "Jeśli po obejrzeniu kalkulatora coś rzuci Ci się w oczy — że czegoś brakuje, "
        "że coś jest niejasne, że tak się tego nie robi — to jest miejsce na to.",
        po=8,
    )
    pole_odpowiedzi(dok, "UWAGI", 8, pytanie=True)

    return dok


if __name__ == "__main__":
    wyjscie = (
        Path(__file__).resolve().parent.parent / "docs" / "Jonasz-formularz-4-pytania.docx"
    )
    wyjscie.parent.mkdir(parents=True, exist_ok=True)
    zbuduj().save(wyjscie)
    print(f"Zapisano: {wyjscie}")
