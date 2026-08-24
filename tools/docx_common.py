"""
Wspólne elementy formularzy .docx do konsultacji ciesielskiej.

Trzyma paletę i funkcje formatujące, żeby kolejne tury pytań wyglądały tak
samo i różniły się wyłącznie treścią.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# --- paleta, spójna ze stroną HTML -------------------------------------------

TUSZ = RGBColor(0x14, 0x20, 0x2C)
TUSZ_SLABY = RGBColor(0x56, 0x63, 0x6F)
AKCENT = RGBColor(0x1F, 0x4E, 0x79)
SYGNAL = RGBColor(0xB0, 0x3D, 0x0D)

TLO_POLA = "F4F6F7"
TLO_POLA_PYTANIE = "FBEEE7"
OBRAMOWANIE = "B3BEC7"
OBRAMOWANIE_PYTANIE = "D9A78E"

# --- narzędzia do formatowania ----------------------------------------------


def cieniuj(komorka, kolor_hex: str) -> None:
    """Nadaje komórce tabeli kolor tła."""
    element = OxmlElement("w:shd")
    element.set(qn("w:val"), "clear")
    element.set(qn("w:fill"), kolor_hex)
    komorka._tc.get_or_add_tcPr().append(element)


def obramuj(komorka, kolor_hex: str, grubosc: int = 8) -> None:
    """Rysuje obramowanie dookoła komórki."""
    borders = OxmlElement("w:tcBorders")
    for krawedz in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{krawedz}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(grubosc))
        el.set(qn("w:color"), kolor_hex)
        borders.append(el)
    komorka._tc.get_or_add_tcPr().append(borders)


def akapit(dok, tekst="", rozmiar=11, bold=False, kolor=TUSZ, kursywa=False,
           przed=0, po=6, wciecie=0.0):
    """Dodaje akapit o zadanym formatowaniu i zwraca go."""
    p = dok.add_paragraph()
    p.paragraph_format.space_before = Pt(przed)
    p.paragraph_format.space_after = Pt(po)
    if wciecie:
        p.paragraph_format.left_indent = Cm(wciecie)
    if tekst:
        run = p.add_run(tekst)
        run.font.size = Pt(rozmiar)
        run.font.bold = bold
        run.font.italic = kursywa
        run.font.color.rgb = kolor
        run.font.name = "Calibri"
    return p


def pole_odpowiedzi(dok, etykieta: str, wiersze: int, pytanie: bool) -> None:
    """Wstawia ramkę, w którą cieśla wpisuje odpowiedź."""
    tabela = dok.add_table(rows=1, cols=1)
    tabela.alignment = WD_TABLE_ALIGNMENT.LEFT
    komorka = tabela.cell(0, 0)
    cieniuj(komorka, TLO_POLA_PYTANIE if pytanie else TLO_POLA)
    obramuj(komorka, OBRAMOWANIE_PYTANIE if pytanie else OBRAMOWANIE)

    pierwszy = komorka.paragraphs[0]
    pierwszy.paragraph_format.space_after = Pt(2)
    run = pierwszy.add_run(etykieta)
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.name = "Calibri"
    run.font.color.rgb = SYGNAL if pytanie else TUSZ_SLABY

    # Puste wiersze dają miejsce na wpisanie odpowiedzi.
    for _ in range(wiersze):
        p = komorka.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.add_run("").font.size = Pt(11)

    dok.add_paragraph().paragraph_format.space_after = Pt(4)


def naglowek(dok, tekst: str, rozmiar=15, kolor=TUSZ, przed=18, po=4, linia=True):
    """Nagłówek sekcji, opcjonalnie z linią pod spodem."""
    p = dok.add_paragraph()
    p.paragraph_format.space_before = Pt(przed)
    p.paragraph_format.space_after = Pt(po)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(tekst)
    run.font.size = Pt(rozmiar)
    run.font.bold = True
    run.font.color.rgb = kolor
    run.font.name = "Calibri"
    if linia:
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), OBRAMOWANIE)
        pbdr.append(bottom)
        p._p.get_or_add_pPr().append(pbdr)
    return p




def nowy_dokument() -> Document:
    """Dokument z ustawionymi marginesami i stylem podstawowym."""
    dok = Document()

    sekcja = dok.sections[0]
    sekcja.top_margin = Cm(2.0)
    sekcja.bottom_margin = Cm(2.0)
    sekcja.left_margin = Cm(2.2)
    sekcja.right_margin = Cm(2.2)

    styl = dok.styles["Normal"]
    styl.font.name = "Calibri"
    styl.font.size = Pt(11)
    styl.font.color.rgb = TUSZ
    return dok


def wypisz_punkty(dok, sekcje) -> None:
    """
    Renderuje sekcje z punktami.

    `sekcje` to lista krotek (tytuł, opis albo None, punkty), gdzie każdy punkt
    to (numer, treść założenia, pytanie albo None).
    """
    for tytul, opis, punkty in sekcje:
        naglowek(dok, tytul)
        if opis:
            akapit(dok, opis, 10, kolor=TUSZ_SLABY, kursywa=True, po=8)

        for numer, zalozenie, pytanie in punkty:
            p = dok.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True

            nr = p.add_run(f"{numer}.  ")
            nr.font.size = Pt(11)
            nr.font.bold = True
            nr.font.color.rgb = SYGNAL if pytanie else AKCENT
            nr.font.name = "Calibri"

            tresc = p.add_run(zalozenie)
            tresc.font.size = Pt(11)
            tresc.font.color.rgb = TUSZ
            tresc.font.name = "Calibri"

            if pytanie:
                pp = dok.add_paragraph()
                pp.paragraph_format.space_before = Pt(0)
                pp.paragraph_format.space_after = Pt(4)
                pp.paragraph_format.left_indent = Cm(0.6)
                pp.paragraph_format.keep_with_next = True
                r = pp.add_run(pytanie)
                r.font.size = Pt(11)
                r.font.bold = True
                r.font.color.rgb = SYGNAL
                r.font.name = "Calibri"
                pole_odpowiedzi(dok, "ODPOWIEDŹ", 3, pytanie=True)
            else:
                pole_odpowiedzi(dok, "UWAGI, JEŚLI COŚ SIĘ NIE ZGADZA", 1, pytanie=False)
