"""
Piąta tura pytań do konsultacji ciesielskiej.

Numeracja punktów jest ciągła z poprzednimi turami i zaczyna się od 112
(tura 1: 1–50, tura 2: 51–73, tura 3: 74–86, tura 4: 87–111). Sekcje mają
tym razem cyfry rzymskie, bo litery skończyły się na turze czwartej.

Ta tura jest inna niż poprzednie. Wcześniej pytaliśmy, jak coś policzyć —
teraz **cała czwarta tura jest wdrożona** i pytamy, czy trafiliśmy. Stąd
dużo pytań zaczyna się od „obejrzyj i powiedz".

Dochodzi też blok, o który nigdy nie pytaliśmy: **wiaty i zadaszenia**.
Ta gałąź powstała na podstawie kart producentów i typowych rozwiązań,
a nie na podstawie praktyki Jonasza — i to jedyna część kalkulatora,
której nie sprawdził cieśla.

Uruchomienie:
    python tools/formularz5-docx.py
"""

from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx.shared import Pt  # noqa: E402

from docx_common import (  # noqa: E402
    SYGNAL,
    TUSZ_SLABY,
    akapit,
    naglowek,
    nowy_dokument,
    pole_odpowiedzi,
    wypisz_punkty,
)

# --- co zmieniły odpowiedzi z czwartej tury --------------------------------

ZMIENIONE = [
    ("Zakładka w kalenicy", "opcja, domyślnie cięcie czołowe", "domyślny sposób zejścia krokwi"),
    ("Koniec krokwi w szczycie", "ucięty na płasko", "ścięty równolegle do krokwi przeciwnej"),
    ("Spięcie zakładki", "nie liczone", "cztery wkręty na parę, długość jak grubość krokwi"),
    ("Okap", "samo cięcie pionowe", "pionowe plus poziome, z miejscem na podbitkę"),
    ("Odsadzka deski podrynnowej", "sztywne 2 cm", "pole do zmiany — to grubość podbitki"),
    ("Długość słupa", "zgadywana połowa wzniesienia", "liczona od podłogi poddasza"),
    ("Szerokość dachu", "tylko rozstaw murłat", "albo rozstaw murłat, albo obrys budynku"),
    ("Okno dachowe", "otwór równy oknu", "luz 1,5 cm po całym obwodzie"),
    ("Wiatrownice", "nie liczone", "dwie łaty po skosie, w zestawieniu"),
    ("Podpisy na modelu", "nie było", "nazwa przy każdym rodzaju elementu"),
    ("Kulawki, kleszcze, miecze", "liczone, ale nierysowane", "widoczne w modelu 3D"),
    ("Pokrycie na modelu", "gładka plama koloru", "rzędy dachówek albo fale blachy, plus gąsior"),
    ("Rzut 2D do druku", "nie było", "rzut z góry z rozstawami, A4 albo A3"),
    ("Deski oparcia w meblach", "leżały płasko jak półki", "w płaszczyźnie oparcia"),
    ("Klej w meblach", "nie liczony", "w zestawieniu, D4 na dwór i D3 do wnętrza"),
    ("Rysunki części mebla", "sama tabela", "każda część w trzech rzutach z wymiarami"),
]

# --- pytania ---------------------------------------------------------------

SEKCJE = [
    (
        "I. Rzut z góry — czy to jest ta kartka",
        "Prosiłeś o „rzut z góry z wymiarami, rozstawami krokwi”. Zrobiliśmy dokładnie "
        "to: obrys z okapami, murłaty, każda krokiew w swoim rzeczywistym rozstawie, "
        "kalenica, a przy kopercie krożyny i kulawki. Otwory pod komin i okno są "
        "zaznaczone. Jest na zakładce „Krokwie”, do wydruku na A4 albo A3.",
        [
            (
                112,
                "Czy na tej kartce czegoś brakuje.",
                "Wydrukuj ją i wyobraź sobie, że rozmierzasz po niej dach. Czego "
                "szukasz wzrokiem i nie znajdujesz?",
            ),
            (
                113,
                "Czy krokwie mają być ponumerowane.",
                "Przy kopercie kulawki są różnej długości i łatwo je pomylić. Czy "
                "numerowanie ich na rzucie pomogłoby, czy tylko zaśmieci rysunek?",
            ),
            (
                114,
                "Czy obok rzutu ma być tabela długości.",
                "Dziś tabela jest osobno, na zakładce „Materiał”. Czy na kartce "
                "zabieranej na budowę ma być razem z rysunkiem, na jednej stronie?",
            ),
        ],
    ),
    (
        "II. Model 3D po poprawkach",
        "Napisałeś: „popracuj nad wizualizacją łat i połączenia w szczycie, ma być "
        "profesjonalnie”. Przy tej okazji wyszło kilka rzeczy, które model rysował "
        "inaczej, niż liczyły wzory — krokiew przechodziła przez środek murłaty "
        "zamiast leżeć na niej, a łaty przebijały przez pokrycie. Poprawione.",
        [
            (
                115,
                "Zamek w kalenicy na modelu.",
                "Obejrzyj szczyt z bliska: krokwie mijają się bokiem, każda ścięta "
                "równolegle do przeciwnej, dolna krawędź jednej dochodzi do górnej "
                "krawędzi drugiej. Czy tak to wygląda na budowie?",
            ),
            (
                116,
                "Zacios na murłacie.",
                "Model rysuje krokiew jako pełną belkę leżącą na murłacie — wrębu nie "
                "pokazuje, bo bryła nie ma jak go wyciąć. Czy to przeszkadza na tyle, "
                "żeby dorobić osobny rysunek samego węzła, czy wystarczy ten, który "
                "jest już na zakładce „Krokwie”?",
            ),
            (
                117,
                "Faktura pokrycia.",
                "Dachówka dostała rzędy w rozstawie łat, blacha trapezowa fale wzdłuż "
                "spadku, kalenica gąsior. Czy z odległości, z której klient ogląda "
                "dach, to wygląda jak materiał, czy jak krata?",
            ),
        ],
    ),
    (
        "III. Meble — sprawdzenie tego, co poprawiliśmy",
        "Napisałeś: „przy oparciach ławek narysowałeś deski w poziomie, gdzie nigdy tak "
        "się nie stosuje — zawsze deska jest montowana prawie pionowo lub pod kątem, "
        "jakim biegnie oparcie”. Zrozumieliśmy to tak, że deska ma LEŻEĆ W PŁASZCZYŹNIE "
        "oparcia, a nie płasko jak półka — i faktycznie trzy meble miały to źle. Ale "
        "zdanie da się przeczytać jeszcze inaczej, więc pytamy wprost.",
        [
            (
                118,
                "Jak mają biec deski oparcia.",
                "Dwie możliwości. (A) Deski biegną WZDŁUŻ ławki, jedna nad drugą, "
                "ustawione w płaszczyźnie oparcia — tak jest teraz. (B) Deski biegną "
                "PIONOWO, od siedziska w górę, jak szczebelki. Która?",
            ),
            (
                119,
                "Od których mebli zacząć.",
                "Napisałeś, że każdy detal trzeba doprecyzować, i masz rację — ale "
                "trzydzieści trzy projekty naraz to za dużo. Wskaż trzy albo cztery, "
                "które robi się najczęściej, a zrobimy je porządnie w pierwszej "
                "kolejności.",
            ),
            (
                120,
                "Kołki i czopy.",
                "Pytaliśmy o złącza, a odpowiedziałeś „licz wkręty i klej ciesielski”. "
                "Klej już liczymy. Czy to znaczy, że kołków i czopów w tych meblach "
                "nie stosujesz w ogóle, czy tylko że nie warto ich liczyć?",
            ),
            (
                121,
                "Rysunki części.",
                "Każda część mebla ma teraz rysunek w trzech rzutach z wymiarami — "
                "z góry, z boku i przekrój. Czy tego brakowało, czy potrzeba jeszcze "
                "rysunku samego złącza: jak dwie części schodzą się w narożu?",
            ),
        ],
    ),
    (
        "IV. Wiaty i zadaszenia — o to nigdy nie pytaliśmy",
        "To druga gałąź kalkulatora, powstała po Twojej uwadze, że „warto zrobić kolejną "
        "gałąź aplikacji: altany i zadaszenia”. Liczy słupy, oczepy, miecze, krokwie, "
        "stopy fundamentowe, rynny i wkręty. Ale wszystkie liczby wzięliśmy z kart "
        "producentów i typowych rozwiązań — nie z praktyki. To jedyna część "
        "kalkulatora, której cieśla nie sprawdził.",
        [
            (
                122,
                "Miecze poprzeczne.",
                "Dajemy parę mieczy w każdej ramie. Czy w praktyce robi się je tylko "
                "w ramach skrajnych, czy faktycznie w każdej?",
            ),
            (
                123,
                "Minimalne spadki pokryć.",
                "Przyjęliśmy: poliwęglan 5°, blacha trapezowa 7°, blachodachówka 12°, "
                "dachówka 25°. Poniżej tych wartości program ostrzega. Czy te progi "
                "zgadzają się z tym, co robisz?",
            ),
            (
                124,
                "Typowe przekroje przy wiacie 5 m.",
                "Słup 140 × 140, oczep 120 × 180, krokiew 80 × 160. Czy tak to wychodzi "
                "u Ciebie, czy któryś z tych przekrojów jest przewymiarowany albo za "
                "słaby?",
            ),
            (
                125,
                "Stopa fundamentowa.",
                "Domyślnie 40 × 40 cm na głębokość 90 cm. Czy to sensowna wartość "
                "startowa dla wiaty na jedno auto?",
            ),
            (
                126,
                "Krokiew na oczepie.",
                "Robisz zacios krokwi na oczepie, tak jak na murłacie, czy podkładasz "
                "klin i mocujesz na płasko?",
            ),
            (
                127,
                "Wymiary gotowych modeli.",
                "W programie jest dwanaście gotowych wiat do wczytania jednym "
                "kliknięciem: wiata na auto 3,5 × 5,5 m, na dwa auta 6 × 5,5 m, "
                "zadaszenie tarasu 3,5 × 6 m, pergola 3,5 × 5 m. Czy tak właśnie "
                "wychodzą te konstrukcje w praktyce?",
            ),
        ],
    ),
]


def zbuduj():
    dok = nowy_dokument()

    naglowek(dok, "Kalkulator ciesielski — piąta tura pytań", 20, przed=0, po=6)
    akapit(
        dok,
        "Cała czwarta tura jest wdrożona i działa na stronie — łącznie z zestawieniem "
        "materiałów przy meblach, które Ci się mieszało z więźbą. Ta tura jest więc "
        "inna niż poprzednie: zamiast pytać, jak coś policzyć, prosimy, żebyś obejrzał "
        "i powiedział, czy trafiliśmy.",
        11,
        po=8,
    )
    akapit(
        dok,
        "Na końcu jest blok o wiatach i zadaszeniach. To jedyna część kalkulatora, "
        "której nigdy nie sprawdzałeś — liczby wzięliśmy z kart producentów, nie "
        "z praktyki.",
        11,
        po=14,
    )

    naglowek(dok, "Co się zmieniło po Twoich odpowiedziach")
    for nazwa, bylo, jest in ZMIENIONE:
        p = dok.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)

        r = p.add_run(f"{nazwa}: ")
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.name = "Calibri"

        r = p.add_run(f"{bylo} → ")
        r.font.size = Pt(10)
        r.font.color.rgb = TUSZ_SLABY
        r.font.name = "Calibri"

        r = p.add_run(jest)
        r.font.size = Pt(10)
        r.font.color.rgb = SYGNAL
        r.font.name = "Calibri"

    akapit(dok, po=10)
    akapit(
        dok,
        "Dwie rzeczy warto wyjaśnić. Wymiary 6,0 / 5,4 / 8,0 cm potraktowaliśmy jako "
        "przykład zaciosu na murłacie, a nie regułę — tak wynikało z Twojej odpowiedzi. "
        "Pisałeś też, że wysłałeś pokazowe zdjęcie zamka: nie dotarło, więc jeśli je "
        "masz, prześlij jeszcze raz.",
        10,
        kolor=TUSZ_SLABY,
        kursywa=True,
        po=6,
    )

    dok.add_page_break()

    wypisz_punkty(dok, SEKCJE)

    # --- pole otwarte ---
    naglowek(dok, "Cokolwiek jeszcze")
    akapit(
        dok,
        "Jak poprzednio: jeśli po obejrzeniu kalkulatora coś rzuci Ci się w oczy — "
        "że czegoś brakuje, że coś jest niejasne, że tak się tego nie robi — to jest "
        "miejsce na to. Poprzednim razem właśnie z tego pola wyszła najważniejsza "
        "poprawka.",
        po=8,
    )
    pole_odpowiedzi(dok, "UWAGI", 8, pytanie=True)

    return dok


if __name__ == "__main__":
    wyjscie = (
        Path(__file__).resolve().parent.parent / "docs" / "Jonasz-formularz-5-pytania.docx"
    )
    wyjscie.parent.mkdir(parents=True, exist_ok=True)
    zbuduj().save(wyjscie)
    print(f"Zapisano: {wyjscie}")
