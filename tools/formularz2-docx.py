"""
Druga tura pytań do konsultacji ciesielskiej.

Pytania wynikają z odpowiedzi Jonasza na pierwszy formularz. Część to
doprecyzowania — powiedział, co jest źle, ale żeby to zaprogramować, trzeba
wiedzieć dokładniej. Reszta to rzeczy, o które w ogóle nie zapytaliśmy.

Numeracja jest ciągła z pierwszym formularzem i zaczyna się od 51, żeby
odwołania do numerów nie kolidowały.

Uruchomienie:
    python tools/formularz2-docx.py
"""

from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx.shared import Cm, Pt  # noqa: E402

from docx_common import (  # noqa: E402
    AKCENT,
    SYGNAL,
    TUSZ,
    TUSZ_SLABY,
    akapit,
    naglowek,
    nowy_dokument,
    pole_odpowiedzi,
    wypisz_punkty,
)

# --- co zmieniły odpowiedzi z pierwszej tury -------------------------------

ZMIENIONE = [
    ("Naddatek na docięcie", "było 5 cm", "jest 10 cm"),
    ("Rzaz piły", "było 4 mm", "jest 5 mm"),
    ("Łaty", "rozstaw + 1 rząd", "rozstaw + łata pod gąsior i na pas okapowy"),
    ("Kleszcze", "wystawały 2 × grubość krokwi", "kończą się na krokwiach"),
    ("Krokiew do murłaty", "dwa kątowniki + 10 wkrętów", "dwa wkręty ciesielskie, kątowniki opcjonalnie"),
    ("Kotwa murłaty", "„kotwa / śruba”", "pręt gwintowany M14–M16 na kotwie chemicznej"),
    ("Kalenica", "gwoździe albo wkręty", "wkręty"),
    ("Impregnat", "liczony zawsze", "opcja — drewno z tartaku bywa impregnowane"),
    ("Wymiary", "w metrach", "w centymetrach, z przełącznikiem na metry"),
]

# --- pytania ---------------------------------------------------------------

SEKCJE = [
    (
        "K. Wysokość słupów — dokończenie punktu 22",
        "Napisałeś: „powinien zapytać o wysokość ścianki kolankowej i murłaty”. Żeby "
        "program policzył z tego długość słupa, musimy wiedzieć dokładnie, co od czego "
        "odmierzać.",
        [
            (
                51,
                "Wysokość ścianki kolankowej.",
                "Od czego się ją mierzy — od stropu poddasza, od podłogi, czy od wieńca "
                "pod spodem? I do czego: do wierzchu wieńca czy do spodu murłaty?",
            ),
            (
                52,
                "Murłata przy ściance kolankowej.",
                "Leży na wieńcu na szczycie ścianki, jest w niego wpuszczona, czy leży "
                "wprost na murze?",
            ),
            (
                53,
                "Na czym stoi słup w więźbie płatwiowej.",
                "Na stropie, na podwalinie, czy bezpośrednio na ściance kolankowej? "
                "Od tego zależy, skąd liczyć jego długość.",
            ),
        ],
    ),
    (
        "L. Kalenica i okap — dokończenie punktów 3 i 6",
        "Napisałeś, że połączenie w kalenicy powinno być do wyboru: czołowe albo zakładka "
        "ciesielska. Przy okapie wspomniałeś o cięciu pionowym i poziomym pod 90°.",
        [
            (
                54,
                "Zakładka ciesielska w kalenicy.",
                "Na jaką głębokość wycina się w każdej krokwi — na pół grubości, czy "
                "inaczej? Jak długie jest samo wycięcie?",
            ),
            (
                55,
                "Długość krokwi przy zakładce.",
                "Czy krokiew przy zakładce jest dłuższa niż przy cięciu czołowym? Jeśli "
                "tak, to o ile — czy zachodzi za oś kalenicy?",
            ),
            (
                56,
                "Cięcie przy desce podrynnowej.",
                "Mówiłeś o cięciu pionowym i poziomym pod 90° do niego. Jak wysokie jest "
                "to poziome cięcie? I czy robi się je zawsze, czy tylko przy niektórych "
                "rynnach?",
            ),
        ],
    ),
    (
        "M. Okna dachowe — dokończenie punktu 49",
        "Napisałeś, że przy oknach robi się rozstaw dopasowany do okien. Chcemy dać taki "
        "tryb rozkładania krokwi, ale potrzebujemy liczb.",
        [
            (
                57,
                "Prześwit pod okno.",
                "Ile luzu zostawiasz z każdej strony okna? Czyli: szerokość okna plus ile "
                "centymetrów daje potrzebny prześwit między krokwiami?",
            ),
            (
                58,
                "Granica, za którą trzeba wymianu.",
                "Od jakiej szerokości okna nie da się już zmieścić w świetle między "
                "krokwiami i trzeba przeciąć krokiew?",
            ),
        ],
    ),
    (
        "N. Murłata a mur — dokończenie punktu 10",
        "Prosiłeś o rubrykę z wizualizacją, ile wysunąć murłatę poza mur. Zanim ją "
        "narysujemy, musimy wiedzieć, jak to bywa w rzeczywistości.",
        [
            (
                59,
                "Położenie murłaty względem muru.",
                "Licuje z zewnętrzną krawędzią muru, jest cofnięta do środka, czy wystaje? "
                "Jeśli cofnięta albo wystająca — o ile zwykle?",
            ),
            (
                60,
                "Program przyjmuje, że rozpiętość mierzy się po zewnętrznych krawędziach "
                "murłat — to potwierdziłeś w punkcie 1.",
                "Czy przy murłacie cofniętej względem muru ta miara nadal jest tą, którą "
                "podaje się z projektu? Czy raczej podaje się wymiar muru, a murłatę "
                "ustawia się osobno?",
            ),
        ],
    ),
    (
        "O. Łączenie krokwi — dokończenie punktu 36",
        "Potwierdziłeś nazwę i sposób spięcia, ale została długość.",
        [
            (
                61,
                "Długość zakładki na styku krokwi. Przyjęliśmy 60 cm.",
                "Ile robi się realnie? Czy zależy to od przekroju krokwi albo od tego, "
                "co jest pod spodem?",
            ),
            (
                62,
                "Spięcie styku. Napisałeś: śrubami z góry i z dołu.",
                "Ile śrub łącznie i jakiej średnicy? Czy przechodzą na wylot z podkładkami "
                "po obu stronach?",
            ),
        ],
    ),
    (
        "P. Wkręty i łaty — dokończenie punktów 45 i 39",
        None,
        [
            (
                63,
                "Wkręty ciesielskie do murłaty. Podałeś 8×220, 8×240 i 10×240 mm.",
                "Od czego zależy wybór — od przekroju krokwi, od kąta dachu, czy po prostu "
                "od tego, co jest pod ręką?",
            ),
            (
                64,
                "Mocowanie samymi wkrętami, bez blach.",
                "Czy przy stromym dachu dokłada się coś przeciw zsuwaniu krokwi, czy dwa "
                "wkręty wystarczają zawsze?",
            ),
            (
                65,
                "Łata pod gąsior i łata na pas okapowy — dołożyliśmy je zgodnie z Twoją uwagą.",
                "Czy są tego samego przekroju co reszta łat, czy grubsze? Pas okapowy to "
                "łata czy raczej deska?",
            ),
        ],
    ),
    (
        "R. Czego w ogóle nie liczymy",
        "To rzeczy, o które nie zapytaliśmy w pierwszej turze, a które schodzą na każdy dach. "
        "Chcemy wiedzieć, czy powinny być w zestawieniu.",
        [
            (
                66,
                "Wiatrownice i deski szczytowe.",
                "Liczyć? Jeśli tak, to jaki przekrój i czy idą po obu szczytach?",
            ),
            (
                67,
                "Deska okapowa i podrynnowa.",
                "Liczyć? Jaki przekrój i na całej długości okapu?",
            ),
            (
                68,
                "Stężenia połaci — wiatrownice ukośne pod krokwiami.",
                "Robi się je na każdym dachu, czy tylko przy większych? Jak je liczyć?",
            ),
            (
                69,
                "Cokolwiek innego, co zawsze schodzi na dach, a nie ma tego u nas.",
                "Co jeszcze zamawiasz razem z więźbą, a o czym w ogóle nie pomyśleliśmy?",
            ),
        ],
    ),
    (
        "S. Strefy śniegowe — Twój pomysł",
        "Zaproponowałeś, żeby kalkulator pytał o województwo i powiat, znał warunki śniegowe "
        "i sugerował przekroje. To duża zmiana, bo dotąd świadomie nie dobieraliśmy przekrojów.",
        [
            (
                70,
                "Zakres tej funkcji.",
                "Ma sugerować konkretne przekroje, czy raczej ostrzegać: „tu jest trzecia "
                "strefa śniegowa, sprawdź to z projektantem”?",
            ),
            (
                71,
                "Kto o to pyta.",
                "Czy klienci realnie pytają o strefy, czy to jest rola projektanta i "
                "kalkulator ma tylko nie przeszkadzać?",
            ),
        ],
    ),
    (
        "T. Rysunki — najczęstsza Twoja uwaga",
        "Wróciłeś do tego cztery razy: więcej obrazków, wizualizacja przy kalenicy, rubryka "
        "z wizualizacją przy murłacie, szata graficzna z motywami drewna. Chcemy to zrobić "
        "dobrze, więc pytamy, co narysować najpierw.",
        [
            (
                72,
                "Kolejność rysunków.",
                "Które trzy miejsca najbardziej potrzebują obrazka? Zacznijmy od nich: "
                "zacios, kalenica, okap z rynną, jętka, słup z płatwią, wymian przy kominie, "
                "rozkład krokwi z góry, coś innego?",
            ),
            (
                73,
                "Rodzaj rysunku.",
                "Rysunek techniczny z wymiarami, jak w projekcie, czy raczej poglądowy "
                "i przestrzenny, żeby klient od razu widział, o co chodzi?",
            ),
        ],
    ),
]


def zbuduj():
    dok = nowy_dokument()

    akapit(dok, "KONSULTACJA CIESIELSKA — DRUGA TURA", 9, bold=True, kolor=AKCENT, po=4)
    akapit(dok, "Kalkulator więźby — pytania po Twoich uwagach", 22, bold=True, po=10)

    akapit(
        dok,
        "Dzięki za pierwszy formularz. Wszystko, co dało się poprawić od razu, jest już "
        "w programie — działająca wersja jest pod adresem podanym przez Marka.",
        po=8,
    )
    akapit(
        dok,
        "Te pytania biorą się z Twoich odpowiedzi. Przy kilku rzeczach powiedziałeś, co jest "
        "źle, ale żeby program to policzył, musimy znać dokładniejsze liczby. Reszta to "
        "sprawy, o które w pierwszej turze w ogóle nie zapytaliśmy.",
        po=8,
    )
    akapit(
        dok,
        "Numeracja jest ciągła z poprzednim formularzem i zaczyna się od 51, więc numery "
        "się nie mylą. Tak jak poprzednio: puste pole czytamy jako „nie mam uwag”.",
        po=14,
    )

    # --- co się zmieniło ---
    naglowek(dok, "Co zmieniły Twoje odpowiedzi", 13, przed=6, po=6)
    akapit(
        dok,
        "Dziewięć rzeczy w programie liczy się teraz inaczej. Dwie z nich zmieniają "
        "zamówienie: kleszcze były o 16 cm za długie na sztukę, a z listy zakupów zniknęło "
        "60 kątowników z 600 wkrętami.",
        10,
        kolor=TUSZ_SLABY,
        po=10,
    )

    for nazwa, bylo, jest in ZMIENIONE:
        p = dok.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.3)

        r = p.add_run(f"{nazwa}: ")
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = TUSZ
        r.font.name = "Calibri"

        r = p.add_run(f"{bylo} → ")
        r.font.size = Pt(10)
        r.font.color.rgb = TUSZ_SLABY
        r.font.name = "Calibri"

        r = p.add_run(jest)
        r.font.size = Pt(10)
        r.font.color.rgb = SYGNAL
        r.font.name = "Calibri"

    akapit(dok, po=10)
    akapit(
        dok,
        "Kąty do dachu kopertowego zostawiliśmy tak, jak były — napisałeś, że to czysta "
        "matematyka, a wzory zgadzają się z tablicami ciesielskimi.",
        10,
        kolor=TUSZ_SLABY,
        kursywa=True,
        po=6,
    )

    dok.add_page_break()

    wypisz_punkty(dok, SEKCJE)

    # --- pole otwarte ---
    naglowek(dok, "Cokolwiek jeszcze")
    akapit(
        dok,
        "Jeśli po obejrzeniu działającego kalkulatora coś rzuci Ci się w oczy — że czegoś "
        "brakuje, że coś jest niejasne, że tak się tego nie robi — to jest miejsce na to.",
        po=8,
    )
    pole_odpowiedzi(dok, "UWAGI", 8, pytanie=True)

    return dok


if __name__ == "__main__":
    wyjscie = (
        Path(__file__).resolve().parent.parent / "docs" / "Jonasz-formularz-2-pytania.docx"
    )
    wyjscie.parent.mkdir(parents=True, exist_ok=True)
    zbuduj().save(wyjscie)
    print(f"Zapisano: {wyjscie}")
